"""
增强版报告导出服务
支持PDF、Word、Excel多种格式导出
"""
import os
import json
import uuid
import hashlib
import datetime
from typing import Dict, List, Optional, Any, BinaryIO
from io import BytesIO

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError as e:
    REPORTLAB_AVAILABLE = False
    print(f"[WARN] reportlab 未安装: {e}")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[WARN] python-docx 未安装，Word导出功能将不可用")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    print("[WARN] pandas 未安装，Excel导出功能将不可用")

from backend.core.config import settings
from backend.models.database import DetectionRecord, Report, User


class EnhancedReportService:
    """增强版报告导出服务"""

    # 特征中文名称映射
    FEATURE_NAMES = {
        "CON_SEM_AI": "语义矛盾度",
        "FIT_TD_AI": "文本-数据一致性",
        "COV_RISK_AI": "风险披露完整性",
        "HIDE_REL_AI": "关联隐藏指数",
        "TONE_ABN_AI": "异常乐观语调",
        "DEN_ABN_AI": "信息密度异常",
        "STR_EVA_AI": "回避表述强度"
    }

    def __init__(self):
        self.template_dir = "backend/templates/reports"
        self.output_dir = settings.REPORT_DIR or "result/reports"
        self._ensure_directories()

    def _ensure_directories(self):
        """确保存储目录存在"""
        os.makedirs(self.template_dir, exist_ok=True)
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_report_file(
        self,
        detection: DetectionRecord,
        report_type: str = "pdf",
        user: Optional[User] = None
    ) -> Dict[str, Any]:
        """
        生成报告文件（支持PDF、Word、Excel）

        Args:
            detection: 检测记录
            report_type: 报告类型 (pdf, word, excel)
            user: 用户信息

        Returns:
            包含文件路径和文件名的字典
        """
        if report_type == "pdf":
            # 使用专业版PDF报告
            from backend.services.professional_report_service import ProfessionalReportService
            service = ProfessionalReportService()
            return service.generate_pdf(detection, user)
        elif report_type == "word":
            return self._generate_word_report(detection, user)
        elif report_type == "excel":
            return self._generate_excel_report(detection, user)
        else:
            raise ValueError(f"不支持的报告类型: {report_type}")

    def _generate_pdf_report(
        self,
        detection: DetectionRecord,
        user: Optional[User] = None
    ) -> Dict[str, Any]:
        """生成PDF报告 - 使用reportlab"""
        if not REPORTLAB_AVAILABLE:
            raise RuntimeError("reportlab 未安装，无法生成PDF")

        # 注册中文字体
        try:
            # 尝试使用微软雅黑
            if os.path.exists('C:/Windows/Fonts/msyh.ttc'):
                pdfmetrics.registerFont(TTFont('ChineseFont', 'C:/Windows/Fonts/msyh.ttc'))
            elif os.path.exists('C:/Windows/Fonts/simhei.ttf'):
                pdfmetrics.registerFont(TTFont('ChineseFont', 'C:/Windows/Fonts/simhei.ttf'))
            else:
                pdfmetrics.registerFont(TTFont('ChineseFont', 'C:/Windows/Fonts/simsun.ttc'))
        except Exception as e:
            print(f'注册中文字体失败: {e}')
            # 使用默认字体，但中文会显示为方框
            pass

        filename = f"{detection.company_name}_{detection.year}_报告_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        file_path = os.path.join(self.output_dir, filename)

        # 创建PDF文档
        doc = SimpleDocTemplate(
            file_path,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )

        # 获取样式
        styles = getSampleStyleSheet()

        # 创建支持中文的自定义样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName='ChineseFont',
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # 居中
            textColor=colors.HexColor('#1f4788')
        )

        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontName='ChineseFont',
            fontSize=16,
            spaceAfter=12,
            spaceBefore=12,
            textColor=colors.HexColor('#1f4788')
        )

        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontName='ChineseFont',
            fontSize=14,
            spaceAfter=10,
            spaceBefore=10,
            textColor=colors.HexColor('#2e5c8a')
        )

        normal_style = ParagraphStyle(
            'NormalChinese',
            parent=styles['Normal'],
            fontName='ChineseFont',
            fontSize=10,
            leading=14
        )

        # 构建PDF内容
        story = []

        # 标题
        story.append(Paragraph("财务舞弊检测报告", title_style))
        story.append(Spacer(1, 0.3 * inch))

        # 企业信息
        story.append(Paragraph("一、企业信息", heading1_style))

        info_data = [
            ['企业名称', detection.company_name],
            ['证券代码', detection.stock_code or "未提供"],
            ['年度', str(detection.year) if detection.year else "未提供"],
            ['检测日期', detection.created_at.strftime("%Y-%m-%d %H:%M:%S")]
        ]

        info_table = Table(info_data, colWidths=[2 * inch, 4 * inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'ChineseFont'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 0.2 * inch))

        # 风险概览
        story.append(Paragraph("二、风险概览", heading1_style))

        risk_level_text = self._get_risk_level_text(detection.risk_level)
        risk_color_hex = '#dc3545' if detection.risk_level == "high" else '#ffc107' if detection.risk_level == "medium" else '#28a745'

        story.append(Paragraph(f"<b>舞弊概率:</b> {detection.fraud_probability:.2%}", normal_style))
        story.append(Paragraph(f"<b>风险等级:</b> <font color='{risk_color_hex}'>{risk_level_text}</font>", normal_style))
        story.append(Paragraph(f"<b>风险评分:</b> {detection.risk_score:.1f}/100", normal_style))
        story.append(Spacer(1, 0.2 * inch))

        # 风险标签
        if detection.risk_labels:
            story.append(Paragraph("三、风险标签", heading1_style))
            for label_info in detection.risk_labels[:8]:  # 最多显示8个
                label = label_info.get("label", "")
                score = label_info.get("score", 0)
                desc = label_info.get("description", "")
                story.append(Paragraph(f"• <b>{label}</b> (评分: {score:.2f})", normal_style))
                if desc:
                    story.append(Paragraph(f"  {desc}", normal_style))
            story.append(Spacer(1, 0.2 * inch))

        # AI特征分析
        if detection.ai_feature_scores:
            story.append(Paragraph("四、AI风险特征分析", heading1_style))
            ai_data = [['特征名称', '风险评分']]
            for feature, score in detection.ai_feature_scores.items():
                if not feature.startswith('_'):
                    # 使用中文名称
                    cn_name = self.FEATURE_NAMES.get(feature, feature)
                    ai_data.append([cn_name, f"{score:.3f}"])

            if len(ai_data) > 1:
                ai_table = Table(ai_data, colWidths=[3 * inch, 2 * inch])
                ai_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'ChineseFont'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
                ]))
                story.append(ai_table)
            story.append(Spacer(1, 0.2 * inch))

        # SHAP特征重要性
        if detection.shap_features:
            story.append(Paragraph("五、SHAP特征重要性", heading1_style))
            shap_data = [['特征名称', 'SHAP值', '影响方向']]
            sorted_features = sorted(detection.shap_features.items(),
                                   key=lambda x: abs(x[1]), reverse=True)[:10]
            for feature, importance in sorted_features:
                # 使用中文名称，避免乱码
                cn_name = self.FEATURE_NAMES.get(feature, feature)
                direction = "推高" if importance > 0 else "抑制"
                shap_data.append([cn_name, f"{importance:+.4f}", direction])

            shap_table = Table(shap_data, colWidths=[2.5 * inch, 1.5 * inch, 1.5 * inch])
            shap_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2e5c8a')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'ChineseFont'),
                ('FONTNAME', (0, 1), (-1, -1), 'ChineseFont'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
            ]))
            story.append(shap_table)
            story.append(Spacer(1, 0.2 * inch))

        # 风险证据
        if detection.risk_evidence_locations:
            story.append(PageBreak())
            story.append(Paragraph("六、风险证据定位", heading1_style))

            for i, evidence in enumerate(detection.risk_evidence_locations[:5], 1):
                feature_name = evidence.get('feature_name', '')
                cn_name = self.FEATURE_NAMES.get(feature_name, feature_name)
                story.append(Paragraph(f"证据 {i}: {cn_name}", heading2_style))
                story.append(Paragraph(f"<b>类别:</b> {evidence.get('category_name', '')}", normal_style))
                story.append(Paragraph(f"<b>位置:</b> {evidence.get('location', '')}", normal_style))
                story.append(Paragraph(f"<b>风险得分:</b> {evidence.get('score', 0):.2f}", normal_style))
                if evidence.get('why_selected'):
                    story.append(Paragraph(f"<b>为什么选择:</b> {evidence['why_selected']}", normal_style))
                if evidence.get('where_is_risk'):
                    story.append(Paragraph(f"<b>风险在哪里:</b> {evidence['where_is_risk']}", normal_style))
                story.append(Spacer(1, 0.1 * inch))

        # 检测建议
        story.append(Paragraph("七、检测建议", heading1_style))
        recommendations = self._generate_recommendations(detection)
        for rec in recommendations:
            story.append(Paragraph(f"• {rec}", normal_style))
        story.append(Spacer(1, 0.2 * inch))

        # 免责声明
        story.append(PageBreak())
        story.append(Paragraph("免责声明", heading2_style))
        disclaimer_text = """本报告基于AI模型分析生成，仅供参考，不构成投资建议或法律意见。
实际风险判断应结合专业审计和实地调查。使用本报告产生的任何后果，本平台不承担责任。"""
        story.append(Paragraph(disclaimer_text, normal_style))

        # 生成PDF
        doc.build(story)

        return {
            "file_path": file_path,
            "filename": filename,
            "file_type": "pdf",
            "file_size": os.path.getsize(file_path)
        }

    def _generate_word_report(
        self,
        detection: DetectionRecord,
        user: Optional[User] = None
    ) -> Dict[str, Any]:
        """生成Word报告"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx 未安装，无法生成Word文档")

        doc = Document()

        # 添加标题
        title = doc.add_heading(f'财务舞弊检测报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 企业信息
        doc.add_heading('一、企业信息', level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_data = [
            ("企业名称", detection.company_name),
            ("证券代码", detection.stock_code or "未提供"),
            ("年度", str(detection.year) if detection.year else "未提供"),
            ("检测日期", detection.created_at.strftime("%Y-%m-%d %H:%M:%S"))
        ]
        for i, (key, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = key
            info_table.rows[i].cells[1].text = str(value)

        # 风险概览
        doc.add_heading('二、风险概览', level=1)
        risk_para = doc.add_paragraph()
        risk_para.add_run(f'舞弊概率: ').bold = True
        risk_run = risk_para.add_run(f'{detection.fraud_probability:.2%}')
        # 根据风险等级设置颜色
        if detection.risk_level == "high":
            risk_run.font.color.rgb = RGBColor(255, 0, 0)
        elif detection.risk_level == "medium":
            risk_run.font.color.rgb = RGBColor(255, 165, 0)
        else:
            risk_run.font.color.rgb = RGBColor(0, 128, 0)

        doc.add_paragraph(f'风险等级: {self._get_risk_level_text(detection.risk_level)}')
        doc.add_paragraph(f'风险评分: {detection.risk_score:.1f}/100')

        # 风险标签
        doc.add_heading('三、风险标签', level=1)
        if detection.risk_labels:
            for label_info in detection.risk_labels:
                label = label_info.get("label", "")
                score = label_info.get("score", 0)
                desc = label_info.get("description", "")

                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f'{label} ').bold = True
                p.add_run(f'(评分: {score:.2f})')
                if desc:
                    doc.add_paragraph(f'  {desc}', style='List Bullet 2')

        # AI特征分析
        doc.add_heading('四、AI风险特征分析', level=1)
        if detection.ai_feature_scores:
            for feature, score in detection.ai_feature_scores.items():
                if not feature.startswith('_'):
                    cn_name = self.FEATURE_NAMES.get(feature, feature)
                    doc.add_paragraph(f'{cn_name} ({feature}): {score:.3f}', style='List Bullet')

        # SHAP特征重要性
        doc.add_heading('五、SHAP特征重要性', level=1)
        if detection.shap_features:
            sorted_features = sorted(detection.shap_features.items(),
                                   key=lambda x: abs(x[1]), reverse=True)[:10]
            for feature, importance in sorted_features:
                cn_name = self.FEATURE_NAMES.get(feature, feature)
                direction = "推高" if importance > 0 else "抑制"
                doc.add_paragraph(f'{cn_name}: {importance:+.4f} ({direction})', style='List Bullet')

        # 风险证据
        if detection.risk_evidence_locations:
            doc.add_heading('六、风险证据定位', level=1)
            for i, evidence in enumerate(detection.risk_evidence_locations[:5], 1):
                feature_name = evidence.get('feature_name', '')
                cn_name = self.FEATURE_NAMES.get(feature_name, feature_name)
                doc.add_heading(f'证据 {i}: {cn_name}', level=2)
                doc.add_paragraph(f'类别: {evidence.get("category_name", "")}')
                doc.add_paragraph(f'位置: {evidence.get("location", "")}')
                doc.add_paragraph(f'风险得分: {evidence.get("score", 0):.2f}')
                if evidence.get("why_selected"):
                    doc.add_paragraph(f'为什么选择: {evidence["why_selected"]}')
                if evidence.get("where_is_risk"):
                    doc.add_paragraph(f'风险位置: {evidence["where_is_risk"]}')

        # 建议
        doc.add_heading('七、检测建议', level=1)
        recommendations = self._generate_recommendations(detection)
        for rec in recommendations:
            doc.add_paragraph(rec, style='List Bullet')

        # 免责声明
        doc.add_page_break()
        disclaimer = doc.add_paragraph()
        disclaimer_run = disclaimer.add_run('免责声明')
        disclaimer_run.bold = True
        disclaimer_run.font.size = Pt(10)
        disclaimer.add_run('\n本报告基于AI模型分析生成，仅供参考，不构成投资建议或法律意见。实际风险判断应结合专业审计和实地调查。')

        # 保存文件
        filename = f"{detection.company_name}_{detection.year}_报告_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)

        return {
            "file_path": file_path,
            "filename": filename,
            "file_type": "word",
            "file_size": os.path.getsize(file_path)
        }

    def _generate_excel_report(
        self,
        detection: DetectionRecord,
        user: Optional[User] = None
    ) -> Dict[str, Any]:
        """生成Excel报告"""
        if not PANDAS_AVAILABLE:
            raise RuntimeError("pandas 未安装，无法生成Excel")

        filename = f"{detection.company_name}_{detection.year}_数据_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        file_path = os.path.join(self.output_dir, filename)

        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            # Sheet 1: 基本信息
            basic_info = {
                '项目': ['企业名称', '证券代码', '年度', '检测日期', '舞弊概率', '风险等级', '风险评分'],
                '值': [
                    detection.company_name,
                    detection.stock_code or '未提供',
                    detection.year or '未提供',
                    detection.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                    f"{detection.fraud_probability:.2%}",
                    self._get_risk_level_text(detection.risk_level),
                    f"{detection.risk_score:.1f}"
                ]
            }
            pd.DataFrame(basic_info).to_excel(writer, sheet_name='基本信息', index=False)

            # Sheet 2: 风险标签
            if detection.risk_labels:
                labels_df = pd.DataFrame([
                    {
                        '风险标签': l.get('label', ''),
                        '类别': l.get('category', ''),
                        '严重程度': l.get('severity', ''),
                        '评分': l.get('score', 0),
                        '描述': l.get('description', ''),
                        '为什么': l.get('why_selected', ''),
                        '风险位置': l.get('where_is_risk', '')
                    }
                    for l in detection.risk_labels
                ])
                labels_df.to_excel(writer, sheet_name='风险标签', index=False)

            # Sheet 3: AI特征分数
            if detection.ai_feature_scores:
                ai_data = []
                for k, v in detection.ai_feature_scores.items():
                    if not k.startswith('_'):
                        cn_name = self.FEATURE_NAMES.get(k, k)
                        ai_data.append({
                            '特征名称': cn_name,
                            '特征代码': k,
                            '分数': v
                        })
                ai_df = pd.DataFrame(ai_data)
                ai_df.to_excel(writer, sheet_name='AI特征分析', index=False)

            # Sheet 4: SHAP特征重要性
            if detection.shap_features:
                shap_data = []
                for k, v in sorted(detection.shap_features.items(),
                                  key=lambda x: abs(x[1]), reverse=True):
                    cn_name = self.FEATURE_NAMES.get(k, k)
                    shap_data.append({
                        '特征名称': cn_name,
                        '特征代码': k,
                        'SHAP值': v,
                        '影响方向': '推高' if v > 0 else '抑制'
                    })
                shap_df = pd.DataFrame(shap_data)
                shap_df.to_excel(writer, sheet_name='SHAP重要性', index=False)

            # Sheet 5: 财务数据
            if detection.financial_data:
                fin_df = pd.DataFrame([
                    {'科目': k, '金额': v}
                    for k, v in detection.financial_data.items()
                ])
                fin_df.to_excel(writer, sheet_name='财务数据', index=False)

            # Sheet 6: 风险证据
            if detection.risk_evidence_locations:
                evidence_data = []
                for e in detection.risk_evidence_locations:
                    feature_name = e.get('feature_name', '')
                    cn_name = self.FEATURE_NAMES.get(feature_name, feature_name)
                    evidence_data.append({
                        '特征名称': cn_name,
                        '特征代码': feature_name,
                        '类别': e.get('category_name', ''),
                        '位置': e.get('location', ''),
                        '评分': e.get('score', 0),
                        'SHAP重要性': e.get('shap_importance', 0),
                        '为什么选择': e.get('why_selected', ''),
                        '风险位置': e.get('where_is_risk', '')
                    })
                evidence_df = pd.DataFrame(evidence_data)
                evidence_df.to_excel(writer, sheet_name='风险证据', index=False)

        return {
            "file_path": file_path,
            "filename": filename,
            "file_type": "excel",
            "file_size": os.path.getsize(file_path)
        }

    def _get_risk_level_text(self, risk_level: str) -> str:
        """获取风险等级文本"""
        descriptions = {
            "low": "低风险 - 舞弊迹象不明显",
            "medium": "中风险 - 存在部分异常信号，需关注",
            "high": "高风险 - 多个舞弊特征显著，建议深入调查"
        }
        return descriptions.get(risk_level, risk_level)

    def _generate_recommendations(self, detection: DetectionRecord) -> List[str]:
        """生成建议"""
        recommendations = []

        if detection.risk_level == "high":
            recommendations.extend([
                "立即停止相关投资决策",
                "聘请专业审计机构进行深入核查",
                "重点关注货币资金、存货、应收账款等关键科目",
                "核查MD&A文本与财务数据的一致性"
            ])
        elif detection.risk_level == "medium":
            recommendations.extend([
                "加强对该企业的财务监控",
                "对比分析同行业可比企业",
                "关注后续财报披露情况",
                "必要时进行专项审计"
            ])
        else:
            recommendations.extend([
                "保持常规关注",
                "定期跟踪财务指标变化",
                "关注行业政策变化影响"
            ])

        return recommendations

    def get_report_as_bytes(self, file_path: str) -> bytes:
        """读取报告文件为字节"""
        with open(file_path, 'rb') as f:
            return f.read()


# 全局服务实例
enhanced_report_service = EnhancedReportService()
