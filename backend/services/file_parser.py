"""
文件解析服务
支持 PDF、Excel、Word 文件解析
提取财务数据和 MD&A 文本
"""
import io
import re
from typing import Dict, Any, List, Optional
import pandas as pd

# 文件解析库
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_SUPPORT = True
except ImportError:
    PYMUPDF_SUPPORT = False

try:
    from openpyxl import load_workbook
    EXCEL_SUPPORT = True
except ImportError:
    EXCEL_SUPPORT = False

try:
    from docx import Document
    WORD_SUPPORT = True
except ImportError:
    WORD_SUPPORT = False


class FileParser:
    """文件解析器"""

    # 财务指标关键词映射
    FINANCIAL_KEYWORDS = {
        "货币资金": ["货币资金", "现金及现金等价物", "Cash and cash equivalents"],
        "短期借款": ["短期借款", "Short-term borrowings", "短期负债"],
        "存货": ["存货", "Inventory", "库存"],
        "营业收入": ["营业收入", "Revenue", "主营业务收入", "营业总收入"],
        "净利润": ["净利润", "Net profit", "归属于母公司股东的净利润", "净利润（净亏损"],
        "总资产": ["总资产", "资产总计", "Total assets", "资产总额"],
        "经营活动现金流净额": ["经营活动现金流", "经营活动产生的现金流量净额", "Operating cash flow"],
        "ROE": ["净资产收益率", "ROE", "Return on equity"],
        "资产负债率": ["资产负债率", "负债比率", "Debt to asset ratio"],
        "营业收入增长率": ["营业收入增长率", "营收增长率", "Revenue growth rate"],
        "净利润增长率": ["净利润增长率", "Net profit growth rate"]
    }

    @staticmethod
    def parse_file(file_content: bytes, filename: str, year: int) -> Dict[str, Any]:
        """
        根据文件类型解析文件

        Args:
            file_content: 文件二进制内容
            filename: 文件名
            year: 年份

        Returns:
            {
                "year": int,
                "financial_data": Dict[str, float],
                "mdna_text": str,
                "parsed_success": bool,
                "parse_errors": List[str]
            }
        """
        file_ext = filename.lower().split('.')[-1]

        parser_map = {
            'pdf': FileParser.parse_pdf,
            'xlsx': FileParser.parse_excel,
            'xls': FileParser.parse_excel,
            'docx': FileParser.parse_word,
            'doc': FileParser.parse_word,
            'txt': FileParser.parse_txt,
            'csv': FileParser.parse_csv
        }

        parser = parser_map.get(file_ext)
        if not parser:
            return {
                "year": year,
                "financial_data": {},
                "mdna_text": "",
                "parsed_success": False,
                "parse_errors": [f"不支持的文件格式: {file_ext}"]
            }

        try:
            result = parser(file_content, year)
            result["filename"] = filename
            return result
        except Exception as e:
            return {
                "year": year,
                "filename": filename,
                "financial_data": {},
                "mdna_text": "",
                "parsed_success": False,
                "parse_errors": [str(e)]
            }

    @staticmethod
    def parse_excel(file_content: bytes, year: int) -> Dict[str, Any]:
        """解析 Excel 文件"""
        if not EXCEL_SUPPORT:
            return {
                "year": year,
                "financial_data": {},
                "mdna_text": "",
                "parsed_success": False,
                "parse_errors": ["未安装 openpyxl，无法解析 Excel 文件"]
            }

        errors = []
        financial_data = {}
        mdna_text = ""

        try:
            # 读取 Excel
            df_dict = pd.read_excel(io.BytesIO(file_content), sheet_name=None)

            # 遍历所有 sheet 寻找财务数据
            for sheet_name, df in df_dict.items():
                # 尝试提取财务数据
                sheet_data = FileParser._extract_from_dataframe(df)
                financial_data.update(sheet_data)

                # 如果找到 MD&A 相关文本
                if "MD" in sheet_name or "管理层" in sheet_name or "讨论" in sheet_name:
                    mdna_text = FileParser._dataframe_to_text(df)

            # 如果没有找到 MD&A，尝试从所有 sheet 中提取文本
            if not mdna_text:
                for sheet_name, df in df_dict.items():
                    text = FileParser._dataframe_to_text(df)
                    if len(text) > 200:  # 至少有一定长度
                        mdna_text = text
                        break

            return {
                "year": year,
                "financial_data": financial_data,
                "mdna_text": mdna_text[:5000],  # 限制长度
                "parsed_success": len(financial_data) > 0 or len(mdna_text) > 0,
                "parse_errors": errors if errors else None
            }

        except Exception as e:
            return {
                "year": year,
                "financial_data": {},
                "mdna_text": "",
                "parsed_success": False,
                "parse_errors": [f"Excel 解析错误: {str(e)}"]
            }

    @staticmethod
    def parse_pdf(file_content: bytes, year: int) -> Dict[str, Any]:
        """解析 PDF 文件"""
        errors = []
        financial_data = {}
        mdna_text = ""

        try:
            # 优先使用 PyMuPDF (更快)
            if PYMUPDF_SUPPORT:
                return FileParser._parse_pdf_with_pymupdf(file_content, year)

            # 备选使用 pdfplumber
            elif PDF_SUPPORT:
                return FileParser._parse_pdf_with_pdfplumber(file_content, year)

            else:
                return {
                    "year": year,
                    "financial_data": {},
                    "mdna_text": "",
                    "parsed_success": False,
                    "parse_errors": ["未安装 PDF 解析库，请安装 pymupdf 或 pdfplumber"]
                }

        except Exception as e:
            return {
                "year": year,
                "financial_data": {},
                "mdna_text": "",
                "parsed_success": False,
                "parse_errors": [f"PDF 解析错误: {str(e)}"]
            }

    @staticmethod
    def _parse_pdf_with_pymupdf(file_content: bytes, year: int) -> Dict[str, Any]:
        """使用 PyMuPDF 解析 PDF"""
        financial_data = {}
        mdna_text = ""
        full_text = ""

        doc = fitz.open(stream=file_content, filetype="pdf")

        # 提取所有文本
        for page in doc:
            full_text += page.get_text()

        doc.close()

        # 提取 MD&A 部分
        mdna_text = FileParser._extract_mdna_from_text(full_text)

        # 尝试从文本中提取财务数据
        financial_data = FileParser._extract_financial_from_text(full_text)

        return {
            "year": year,
            "financial_data": financial_data,
            "mdna_text": mdna_text[:5000],
            "parsed_success": len(financial_data) > 0 or len(mdna_text) > 0,
            "parse_errors": None
        }

    @staticmethod
    def _parse_pdf_with_pdfplumber(file_content: bytes, year: int) -> Dict[str, Any]:
        """使用 pdfplumber 解析 PDF"""
        financial_data = {}
        mdna_text = ""
        full_text = ""

        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n"

                # 尝试从表格提取数据
                tables = page.extract_tables()
                for table in tables:
                    table_data = FileParser._extract_from_table(table)
                    financial_data.update(table_data)

        # 提取 MD&A
        mdna_text = FileParser._extract_mdna_from_text(full_text)

        # 如果表格没提取到，尝试从文本提取
        if not financial_data:
            financial_data = FileParser._extract_financial_from_text(full_text)

        return {
            "year": year,
            "financial_data": financial_data,
            "mdna_text": mdna_text[:5000],
            "parsed_success": len(financial_data) > 0 or len(mdna_text) > 0,
            "parse_errors": None
        }

    @staticmethod
    def parse_word(file_content: bytes, year: int) -> Dict[str, Any]:
        """解析 Word 文件"""
        if not WORD_SUPPORT:
            return {
                "year": year,
                "financial_data": {},
                "mdna_text": "",
                "parsed_success": False,
                "parse_errors": ["未安装 python-docx，无法解析 Word 文件"]
            }

        try:
            doc = Document(io.BytesIO(file_content))

            # 提取所有文本
            full_text = ""
            for para in doc.paragraphs:
                full_text += para.text + "\n"

            # 提取表格数据
            financial_data = {}
            for table in doc.tables:
                table_data = FileParser._extract_from_docx_table(table)
                financial_data.update(table_data)

            # 提取 MD&A
            mdna_text = FileParser._extract_mdna_from_text(full_text)

            # 如果没从表格提取到，尝试从文本提取
            if not financial_data:
                financial_data = FileParser._extract_financial_from_text(full_text)

            return {
                "year": year,
                "financial_data": financial_data,
                "mdna_text": mdna_text[:5000],
                "parsed_success": len(financial_data) > 0 or len(mdna_text) > 0,
                "parse_errors": None
            }

        except Exception as e:
            return {
                "year": year,
                "financial_data": {},
                "mdna_text": "",
                "parsed_success": False,
                "parse_errors": [f"Word 解析错误: {str(e)}"]
            }

    @staticmethod
    def parse_txt(file_content: bytes, year: int) -> Dict[str, Any]:
        """解析纯文本文件"""
        try:
            # 尝试多种编码方式解码
            encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'utf-16', 'latin-1']
            full_text = ""

            for encoding in encodings:
                try:
                    full_text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if not full_text:
                return {
                    "year": year,
                    "financial_data": {},
                    "mdna_text": "",
                    "parsed_success": False,
                    "parse_errors": ["无法识别文件编码格式"]
                }

            # 尝试从文本中提取 MD&A
            mdna_text = FileParser._extract_mdna_from_text(full_text)

            # 如果没能提取到特定 MD&A 部分，使用全文
            if not mdna_text or len(mdna_text) < 100:
                mdna_text = full_text

            # 尝试从文本中提取财务数据
            financial_data = FileParser._extract_financial_from_text(full_text)

            return {
                "year": year,
                "financial_data": financial_data,
                "mdna_text": mdna_text[:5000],
                "parsed_success": True,
                "parse_errors": None
            }

        except Exception as e:
            return {
                "year": year,
                "financial_data": {},
                "mdna_text": "",
                "parsed_success": False,
                "parse_errors": [f"文本文件解析错误: {str(e)}"]
            }

    @staticmethod
    def parse_csv(file_content: bytes, year: int) -> Dict[str, Any]:
        """解析 CSV 文件"""
        try:
            # 尝试多种编码方式
            encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'utf-16']
            df = None

            for encoding in encodings:
                try:
                    df = pd.read_csv(io.BytesIO(file_content), encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
                except Exception:
                    continue

            if df is None:
                return {
                    "year": year,
                    "financial_data": {},
                    "mdna_text": "",
                    "parsed_success": False,
                    "parse_errors": ["无法解析 CSV 文件编码"]
                }

            # 从 DataFrame 中提取财务数据
            financial_data = FileParser._extract_from_dataframe(df)

            # 将 DataFrame 转换为文本作为 MD&A 内容
            mdna_text = FileParser._dataframe_to_text(df)

            return {
                "year": year,
                "financial_data": financial_data,
                "mdna_text": mdna_text[:5000],
                "parsed_success": len(financial_data) > 0 or len(mdna_text) > 0,
                "parse_errors": None
            }

        except Exception as e:
            return {
                "year": year,
                "financial_data": {},
                "mdna_text": "",
                "parsed_success": False,
                "parse_errors": [f"CSV 解析错误: {str(e)}"]
            }

    @staticmethod
    def _extract_from_dataframe(df: pd.DataFrame) -> Dict[str, float]:
        """从 DataFrame 提取财务数据"""
        result = {}

        # 将 DataFrame 转为字符串搜索
        df_str = df.to_string()

        for field, keywords in FileParser.FINANCIAL_KEYWORDS.items():
            for keyword in keywords:
                # 在 DataFrame 中搜索关键词
                mask = df.astype(str).apply(lambda x: x.str.contains(keyword, na=False))
                if mask.any().any():
                    # 找到包含关键词的行，尝试提取数值
                    for idx in df[mask.any(axis=1)].index:
                        row = df.loc[idx]
                        # 尝试从行中提取数值
                        for val in row:
                            if isinstance(val, (int, float)) and val != 0:
                                result[field] = float(val)
                                break
                        if field in result:
                            break
                if field in result:
                    break

        return result

    @staticmethod
    def _extract_from_table(table: List[List]) -> Dict[str, float]:
        """从表格数据提取财务数据"""
        result = {}

        for row in table:
            row_text = " ".join([str(cell) for cell in row if cell])

            for field, keywords in FileParser.FINANCIAL_KEYWORDS.items():
                for keyword in keywords:
                    if keyword in row_text:
                        # 尝试从行中提取数值
                        for cell in row:
                            try:
                                # 清理字符串中的逗号和单位
                                if isinstance(cell, str):
                                    cell = cell.replace(",", "").replace("¥", "").replace("元", "")
                                    cell = cell.replace("亿元", "").replace("万", "")

                                val = float(cell)
                                if val != 0:
                                    # 如果是亿元单位，转换为元
                                    if "亿元" in str(cell) or (val < 10000 and val > 100):
                                        val = val * 100000000
                                    # 如果是万元单位，转换为元
                                    elif "万" in str(cell):
                                        val = val * 10000

                                    result[field] = val
                                    break
                            except:
                                continue
                        if field in result:
                            break

        return result

    @staticmethod
    def _extract_from_docx_table(table) -> Dict[str, float]:
        """从 docx 表格提取数据"""
        result = {}

        for row in table.rows:
            row_text = " ".join([cell.text for cell in row.cells])

            for field, keywords in FileParser.FINANCIAL_KEYWORDS.items():
                for keyword in keywords:
                    if keyword in row_text:
                        # 尝试提取数值
                        for cell in row.cells:
                            try:
                                text = cell.text.replace(",", "").replace("¥", "").replace("元", "")
                                text = text.replace("亿元", "").replace("万", "")
                                val = float(text)
                                if val != 0:
                                    if "亿元" in cell.text or (val < 10000 and val > 100):
                                        val = val * 100000000
                                    elif "万" in cell.text:
                                        val = val * 10000
                                    result[field] = val
                                    break
                            except:
                                continue
                        if field in result:
                            break

        return result

    @staticmethod
    def _extract_mdna_from_text(text: str) -> str:
        """从文本中提取 MD&A 部分"""
        # 常见的 MD&A 章节标题
        mdna_patterns = [
            r"管理层讨论与分析.*?(?=第五节|第六节|第六章|$)",
            r"MD&A.*?(?=\n\n|\r\n\r\n|$)",
            r"第四节\s*管理层讨论与分析.*?(?=第五节|第六节|$)",
            r"经营情况讨论与分析.*?(?=\n\n|$)",
        ]

        for pattern in mdna_patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                return match.group(0).strip()

        # 如果没找到特定章节，返回前5000字符
        return text[:5000]

    @staticmethod
    def _extract_financial_from_text(text: str) -> Dict[str, float]:
        """从文本中提取财务数据"""
        result = {}

        for field, keywords in FileParser.FINANCIAL_KEYWORDS.items():
            for keyword in keywords:
                # 搜索关键词后的数值
                pattern = rf"{re.escape(keyword)}.*?([\d,]+\.?\d*)"
                matches = re.findall(pattern, text, re.IGNORECASE)

                for match in matches:
                    try:
                        val = float(match.replace(",", ""))
                        if val > 0:
                            result[field] = val
                            break
                    except:
                        continue

                if field in result:
                    break

        return result

    @staticmethod
    def _dataframe_to_text(df: pd.DataFrame) -> str:
        """将 DataFrame 转为文本"""
        return df.to_string(index=False)


# 便捷函数
def parse_financial_file(file_content: bytes, filename: str, year: int) -> Dict[str, Any]:
    """
    解析财务文件的便捷函数

    Args:
        file_content: 文件二进制内容
        filename: 文件名
        year: 年份

    Returns:
        解析结果字典
    """
    return FileParser.parse_file(file_content, filename, year)


def merge_parsed_results(results: List[Dict[str, Any]], year: int) -> Dict[str, Any]:
    """
    合并同一年份多个文件的解析结果
    用于处理结构化文档（Excel/财务表）和非结构化文档（PDF/MD&A）分开上传的情况

    Args:
        results: 多个文件的解析结果列表
        year: 年份

    Returns:
        合并后的解析结果
    """
    merged_financial_data = {}
    merged_mdna_texts = []
    all_errors = []
    parsed_file_names = []

    for result in results:
        if not result.get("parsed_success"):
            all_errors.extend(result.get("parse_errors", []))
            continue

        # 记录文件名
        if result.get("filename"):
            parsed_file_names.append(result["filename"])

        # 合并财务数据（结构化数据优先，后覆盖前）
        financial_data = result.get("financial_data", {})
        for key, value in financial_data.items():
            if value and value != 0:  # 只保留非空值
                merged_financial_data[key] = value

        # 收集 MD&A 文本（非结构化数据）
        mdna_text = result.get("mdna_text", "")
        if mdna_text and len(mdna_text) > 50:  # 至少有意义的文本长度
            merged_mdna_texts.append({
                "source": result.get("filename", "unknown"),
                "text": mdna_text
            })

        # 收集错误
        if result.get("parse_errors"):
            all_errors.extend(result["parse_errors"])

    # 合并 MD&A 文本，添加分隔标识
    final_mdna_text = ""
    if len(merged_mdna_texts) == 1:
        final_mdna_text = merged_mdna_texts[0]["text"]
    elif len(merged_mdna_texts) > 1:
        # 多个 MD&A 文本，用分隔符拼接
        parts = []
        for item in merged_mdna_texts:
            parts.append(f"【来源: {item['source']}】\n{item['text']}")
        final_mdna_text = "\n\n---\n\n".join(parts)

    # 判断是否成功：至少要有财务数据或 MD&A 文本
    has_data = len(merged_financial_data) > 0 or len(final_mdna_text) > 0

    return {
        "year": year,
        "financial_data": merged_financial_data,
        "mdna_text": final_mdna_text,
        "parsed_success": has_data,
        "parse_errors": all_errors if all_errors else None,
        "source_files": parsed_file_names,
        "data_sources": {
            "financial_data_sources": [r.get("filename") for r in results if r.get("financial_data")],
            "mdna_text_sources": [item["source"] for item in merged_mdna_texts]
        }
    }
